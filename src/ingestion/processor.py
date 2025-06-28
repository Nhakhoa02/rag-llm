"""
Document processing module for handling different file types.
"""

import os
import io
from typing import List, Dict, Any, Optional, Union, Tuple
from pathlib import Path
import pandas as pd
from PIL import Image
import fitz  # PyMuPDF. Ensure you have only 'PyMuPDF' installed, not the unrelated 'fitz' package from PyPI.
from docx import Document as DocxDocument
import json
import xml.etree.ElementTree as ET
from pypdf import PdfReader
import google.generativeai as genai

from ..models.base import BaseDocument, DataType, ProcessingStatus
from ..models.csv_index import CSVIndex, CSVIndexDocument
from ..utils.logging import LoggerMixin
from ..utils.validation import data_validator
from ..utils.metrics import monitor_function


class DocumentProcessor(LoggerMixin):
    """Base document processor for handling different file types."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_text,
            '.json': self._process_json,
            '.xml': self._process_xml,
            '.csv': self._process_csv,
            '.xlsx': self._process_excel,
        }
    
    @monitor_function("document_processor", "process_document", "document")
    def process_document(self, file_path: Union[str, Path], 
                        metadata: Optional[Dict[str, Any]] = None) -> BaseDocument:
        """
        Process a document and return a BaseDocument object.
        
        Args:
            file_path: Path to the document file
            metadata: Additional metadata for the document
            
        Returns:
            BaseDocument object with processed content
        """
        try:
            file_path = Path(file_path)
            
            # Validate file
            is_valid, error_msg = data_validator.validate_file_type(file_path)
            if not is_valid:
                raise ValueError(f"File validation failed: {error_msg}")
            
            is_valid, error_msg = data_validator.validate_file_size(file_path)
            if not is_valid:
                raise ValueError(f"File size validation failed: {error_msg}")
            
            # Determine file type
            extension = file_path.suffix.lower()
            if extension not in self.supported_extensions:
                raise ValueError(f"Unsupported file type: {extension}")
            
            # Process file
            content = self.supported_extensions[extension](file_path)
            
            # Create document object
            doc = BaseDocument(
                type=DataType.DOCUMENT,
                content=content,
                metadata=metadata or {}
            )
            
            # Add file metadata
            doc.update_metadata("file_path", str(file_path))
            doc.update_metadata("file_size", file_path.stat().st_size)
            doc.update_metadata("file_extension", extension)
            doc.update_metadata("file_hash", data_validator.calculate_file_hash(file_path))
            
            doc.status = ProcessingStatus.COMPLETED
            self.logger.info("Document processed successfully", document_id=doc.id, file_path=str(file_path))
            
            return doc
            
        except Exception as e:
            self.logger.error("Document processing failed", error=str(e), file_path=str(file_path))
            raise
    
    def _process_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            # Try PyMuPDF first (better text extraction)
            doc = fitz.open(str(file_path))
            text = ""
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                if page_text:
                    text += page_text + "\n"
                else:
                    self.logger.warning(f"No text extracted from page {page_num + 1}")
            doc.close()
            
            text = text.strip()
            if not text:
                self.logger.warning("No text extracted from PDF using PyMuPDF, trying PyPDF fallback")
                raise Exception("No text extracted")
            
            self.logger.info(f"Successfully extracted {len(text)} characters from PDF using PyMuPDF")
            return text
            
        except Exception as e:
            self.logger.warning(f"PyMuPDF failed: {e}, trying PyPDF fallback")
            try:
                # Fallback to PyPDF
                with open(file_path, 'rb') as file:
                    reader = PdfReader(file)
                    text = ""
                    for page_num, page in enumerate(reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                        else:
                            self.logger.warning(f"No text extracted from page {page_num + 1} using PyPDF")
                    
                    text = text.strip()
                    if not text:
                        self.logger.error("No text could be extracted from PDF using either method")
                        return "PDF document (no text content extracted)"
                    
                    self.logger.info(f"Successfully extracted {len(text)} characters from PDF using PyPDF")
                    return text
                    
            except Exception as fallback_error:
                self.logger.error(f"Both PDF extraction methods failed: {fallback_error}")
                return "PDF document (text extraction failed)"
    
    def _process_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        doc = DocxDocument(str(file_path))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    
    def _process_text(self, file_path: Path) -> str:
        """Read text file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    
    def _process_json(self, file_path: Path) -> str:
        """Process JSON file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
            return json.dumps(data, indent=2)
    
    def _process_xml(self, file_path: Path) -> str:
        """Process XML file."""
        tree = ET.parse(file_path)
        root = tree.getroot()
        return ET.tostring(root, encoding='unicode')
    
    def _process_csv(self, file_path: Path) -> str:
        """Process CSV file."""
        df = pd.read_csv(file_path)
        return df.to_string(index=False)
    
    def _process_excel(self, file_path: Path) -> str:
        """Process Excel file."""
        df = pd.read_excel(file_path)
        return df.to_string(index=False)


class ImageProcessor(LoggerMixin):
    """Processor for image files with OCR and feature extraction."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'}
    
    @monitor_function("image_processor", "process_image", "image")
    def process_image(self, file_path: Union[str, Path], metadata: Optional[Dict[str, Any]] = None) -> BaseDocument:
        """
        Process an image and generate a caption using Gemini 2.5 Flash.
        """
        try:
            file_path = Path(file_path)
            extension = file_path.suffix.lower()
            image = Image.open(file_path).convert("RGB")
            # Use Gemini Vision for caption
            caption = self.get_gemini_image_caption(str(file_path))
            self.logger.info("Gemini Vision caption generated for image", caption=caption)
            features = self._extract_image_features(image)
            doc = BaseDocument(
                type=DataType.IMAGE,
                content=caption,
                metadata=metadata or {}
            )
            
            # Add image metadata
            doc.update_metadata("file_path", str(file_path))
            doc.update_metadata("file_size", file_path.stat().st_size)
            doc.update_metadata("file_extension", extension)
            doc.update_metadata("image_features", features)
            doc.update_metadata("file_hash", data_validator.calculate_file_hash(file_path))
            
            doc.status = ProcessingStatus.COMPLETED
            self.logger.info("Image processed successfully", document_id=doc.id, file_path=str(file_path))
            
            return doc
            
        except Exception as e:
            self.logger.error("Image processing failed", error=str(e), file_path=str(file_path))
            raise
    
    def get_gemini_image_caption(self, image_path: str) -> str:
        """Generate image caption using Gemini Vision."""
        try:
            # Configure Gemini
            genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
            model = genai.GenerativeModel('gemini-2.0-flash-exp')
            
            # Load and process image
            image = Image.open(image_path)
            
            # Generate caption
            response = model.generate_content([
                "Describe this image in detail, focusing on the main subjects, actions, and context. Be specific about what you see.",
                image
            ])
            
            # Extract text from response
            if hasattr(response, "text") and response.text:
                return response.text.strip()
            
            # Fallback to parts if text doesn't work
            if hasattr(response, "parts") and response.parts:
                text_parts = []
                for part in response.parts:
                    if hasattr(part, "text") and part.text:
                        text_parts.append(part.text)
                return " ".join(text_parts)
        
        except Exception as e:
            self.logger.warning(f"Gemini Vision caption generation failed: {e}")
            return f"Image file: {os.path.basename(image_path)}"
        
        # Fallback to candidates if parts don't work
        if hasattr(response, "candidates") and response.candidates:
            text_parts = []
            for candidate in response.candidates:
                if hasattr(candidate, "content") and hasattr(candidate.content, "parts"):
                    for part in candidate.content.parts:
                        if hasattr(part, "text") and part.text:
                            text_parts.append(part.text)
            return " ".join(text_parts)
        
        return ""

    def _extract_image_features(self, image: Image.Image) -> Dict[str, Any]:
        try:
            features = {
                "width": image.width,
                "height": image.height,
                "aspect_ratio": image.width / image.height,
                "mode": image.mode,
                "format": image.format,
            }
            return features
        except Exception as e:
            self.logger.warning("Feature extraction failed", error=str(e))
            return {"width": image.width, "height": image.height, "mode": image.mode}


class TabularProcessor(LoggerMixin):
    """Processor for tabular data (CSV, Excel, etc.) with CSV indexing support."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = {'.csv', '.xlsx', '.xls', '.parquet', '.json'}
    
    @monitor_function("tabular_processor", "process_tabular", "tabular")
    def process_tabular(self, file_path: Union[str, Path], 
                       metadata: Optional[Dict[str, Any]] = None) -> BaseDocument:
        """
        Process tabular data and extract structured information.
        
        Args:
            file_path: Path to the tabular file
            metadata: Additional metadata for the data
            
        Returns:
            BaseDocument object with processed content
        """
        try:
            file_path = Path(file_path)
            
            # Validate file
            extension = file_path.suffix.lower()
            if extension not in self.supported_extensions:
                raise ValueError(f"Unsupported tabular format: {extension}")
            
            # Load data
            df = self._load_tabular_data(file_path)
            
            # Extract structured information
            structure_info = self._extract_structure_info(df)
            
            # Convert to text representation
            text_content = self._convert_to_text(df)
            
            # Create document object
            doc = BaseDocument(
                type=DataType.TABULAR,
                content=text_content,
                metadata=metadata or {}
            )
            
            # Add tabular metadata
            doc.update_metadata("file_path", str(file_path))
            doc.update_metadata("row_count", len(df))
            doc.update_metadata("column_count", len(df.columns))
            doc.update_metadata("columns", df.columns.tolist())
            doc.update_metadata("data_types", df.dtypes.astype(str).to_dict())
            doc.update_metadata("structure_info", structure_info)
            doc.update_metadata("file_hash", data_validator.calculate_file_hash(file_path))
            
            # For CSV files, create and store CSV index
            if extension == '.csv':
                csv_index = self._create_csv_index(df, doc.id, file_path.name)
                doc.update_metadata("csv_index", csv_index.to_dict())
                self.logger.info(f"Created CSV index for {file_path.name}")
            
            doc.status = ProcessingStatus.COMPLETED
            self.logger.info("Tabular data processed successfully", document_id=doc.id, file_path=str(file_path))
            
            return doc
            
        except Exception as e:
            self.logger.error("Tabular processing failed", error=str(e), file_path=str(file_path))
            raise
    
    def _load_tabular_data(self, file_path: Path) -> pd.DataFrame:
        """Load tabular data from file."""
        extension = file_path.suffix.lower()
        
        if extension == '.csv':
            return pd.read_csv(file_path)
        elif extension in ['.xlsx', '.xls']:
            return pd.read_excel(file_path)
        elif extension == '.parquet':
            return pd.read_parquet(file_path)
        elif extension == '.json':
            return pd.read_json(file_path)
        else:
            raise ValueError(f"Unsupported tabular format: {extension}")
    
    def _extract_structure_info(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Extract structural information from dataframe."""
        info = {
            "shape": df.shape,
            "columns": df.columns.tolist(),
            "data_types": df.dtypes.astype(str).to_dict(),
            "missing_values": df.isnull().sum().to_dict(),
            "unique_counts": {col: df[col].nunique() for col in df.columns},
        }
        
        # Add statistical information for numeric columns
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            info["numeric_stats"] = df[numeric_cols].describe().to_dict()
        
        return info
    
    def _create_csv_index(self, df: pd.DataFrame, csv_file_id: str, filename: str) -> CSVIndex:
        """Create a CSV index from the first and second rows of the dataframe."""
        try:
            # Extract first row (column headers)
            column_headers = df.columns.tolist()
            
            # Extract second row (sample data) if available
            sample_data = None
            if len(df) > 0:
                second_row = df.iloc[0].tolist()  # First data row (index 0)
                sample_data = [str(val) if pd.notna(val) else "" for val in second_row]
            
            # Infer data types from sample
            inferred_types = {}
            if sample_data:
                for header, value in zip(column_headers, sample_data):
                    if value:
                        try:
                            # Try to infer type
                            if value.replace('.', '').replace('-', '').isdigit():
                                if '.' in value:
                                    inferred_types[header] = 'float'
                                else:
                                    inferred_types[header] = 'int'
                            else:
                                inferred_types[header] = 'string'
                        except:
                            inferred_types[header] = 'string'
            
            # Create CSV index
            csv_index = CSVIndex(
                csv_file_id=csv_file_id,
                csv_filename=filename,
                column_headers=column_headers,
                sample_data=sample_data,
                total_rows=len(df),
                total_columns=len(df.columns),
                inferred_types=inferred_types
            )
            
            # Generate index content
            csv_index.update_index_content()
            
            return csv_index
            
        except Exception as e:
            self.logger.error(f"Failed to create CSV index: {e}")
            # Return a minimal index
            return CSVIndex(
                csv_file_id=csv_file_id,
                csv_filename=filename,
                column_headers=df.columns.tolist(),
                total_rows=len(df),
                total_columns=len(df.columns)
            )
    
    def create_csv_index_document(self, csv_index: CSVIndex) -> CSVIndexDocument:
        """Create a CSV index document for vector storage."""
        if csv_index is None:
            raise ValueError("CSV index cannot be None")
        return CSVIndexDocument(csv_index=csv_index)
    
    def _convert_to_text(self, df: pd.DataFrame) -> str:
        """Convert dataframe to text representation that's AI-friendly."""
        # Create a more structured, readable format for AI processing
        text_parts = []
        
        # Add dataset overview
        text_parts.append(f"Dataset Overview:")
        text_parts.append(f"- Total rows: {len(df)}")
        text_parts.append(f"- Total columns: {len(df.columns)}")
        text_parts.append(f"- Column names: {', '.join(df.columns.tolist())}")
        text_parts.append("")
        
        # Add data types information
        text_parts.append("Column Data Types:")
        for col, dtype in df.dtypes.items():
            text_parts.append(f"- {col}: {dtype}")
        text_parts.append("")
        
        # Add missing values summary
        missing_values = df.isnull().sum()
        if missing_values.sum() > 0:
            text_parts.append("Missing Values Summary:")
            for col, missing_count in missing_values.items():
                if missing_count > 0:
                    percentage = (missing_count / len(df)) * 100
                    text_parts.append(f"- {col}: {missing_count} missing values ({percentage:.1f}%)")
            text_parts.append("")
        
        # Add sample data in a more readable format
        text_parts.append("Sample Data:")
        
        # Limit the number of rows for readability
        if len(df) > 50:
            # Show first 25 and last 25 rows
            sample_df = pd.concat([df.head(25), df.tail(25)], ignore_index=True)
            text_parts.append(f"(Showing first 25 and last 25 rows out of {len(df)} total rows)")
        else:
            sample_df = df
        
        # Convert to a more readable format
        for i in range(len(sample_df)):
            row = sample_df.iloc[i]
            row_text = f"Row {i + 1}: "
            row_data = []
            for col in df.columns:
                value = row[col]
                # Handle different data types
                if pd.isna(value):
                    row_data.append(f"{col}=NULL")
                elif isinstance(value, (int, float)):
                    row_data.append(f"{col}={value}")
                else:
                    # Truncate long string values
                    str_value = str(value)
                    if len(str_value) > 50:
                        str_value = str_value[:47] + "..."
                    row_data.append(f"{col}='{str_value}'")
            row_text += ", ".join(row_data)
            text_parts.append(row_text)
        
        # Add summary statistics for numeric columns
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            text_parts.append("")
            text_parts.append("Numeric Column Statistics:")
            for col in numeric_cols:
                stats = df[col].describe()
                text_parts.append(f"- {col}:")
                text_parts.append(f"  * Count: {stats['count']:.0f}")
                text_parts.append(f"  * Mean: {stats['mean']:.2f}")
                text_parts.append(f"  * Min: {stats['min']:.2f}")
                text_parts.append(f"  * Max: {stats['max']:.2f}")
                text_parts.append(f"  * Std: {stats['std']:.2f}")
        
        return "\n".join(text_parts) 