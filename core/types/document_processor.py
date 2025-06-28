"""
Document processing module for handling different file types.
"""

import os
import io
from typing import List, Dict, Any, Optional, Union, Tuple
from pathlib import Path
import pandas as pd
from PIL import Image
import fitz  # PyMuPDF. Ensure you have only 'PyMuPDF' installed, not the unrelated 'fitz' package from PyPI.
from docx import Document as DocxDocument
import json
import xml.etree.ElementTree as ET
from pypdf import PdfReader
import google.generativeai as genai

from core.models.base import BaseDocument, DataType, ProcessingStatus
from core.utils.logging import LoggerMixin
from core.utils.validation import data_validator
from core.utils.metrics import monitor_function


class DocumentProcessor(LoggerMixin):
    """Base document processor for handling different file types."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_text,
            '.json': self._process_json,
            '.xml': self._process_xml,
            '.csv': self._process_csv,
            '.xlsx': self._process_excel,
        }
    
    @monitor_function("document_processor", "process_document", "document")
    def process_document(self, file_path: Union[str, Path], 
                        metadata: Optional[Dict[str, Any]] = None) -> BaseDocument:
        """
        Process a document and return a BaseDocument object.
        
        Args:
            file_path: Path to the document file
            metadata: Additional metadata for the document
            
        Returns:
            BaseDocument object with processed content
        """
        try:
            file_path = Path(file_path)
            
            # Validate file
            is_valid, error_msg = data_validator.validate_file_type(file_path)
            if not is_valid:
                raise ValueError(f"File validation failed: {error_msg}")
            
            is_valid, error_msg = data_validator.validate_file_size(file_path)
            if not is_valid:
                raise ValueError(f"File size validation failed: {error_msg}")
            
            # Determine file type
            extension = file_path.suffix.lower()
            if extension not in self.supported_extensions:
                raise ValueError(f"Unsupported file type: {extension}")
            
            # Process file
            content = self.supported_extensions[extension](file_path)
            
            # Create document object
            doc = BaseDocument(
                type=DataType.DOCUMENT,
                content=content,
                metadata=metadata or {}
            )
            
            # Add file metadata
            doc.update_metadata("file_path", str(file_path))
            doc.update_metadata("file_size", file_path.stat().st_size)
            doc.update_metadata("file_extension", extension)
            doc.update_metadata("file_hash", data_validator.calculate_file_hash(file_path))
            
            doc.status = ProcessingStatus.COMPLETED
            self.logger.info("Document processed successfully", document_id=doc.id, file_path=str(file_path))
            
            return doc
            
        except Exception as e:
            self.logger.error("Document processing failed", error=str(e), file_path=str(file_path))
            raise
    
    def _process_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            # Try PyMuPDF first (better text extraction)
            doc = fitz.open(str(file_path))
            text = ""
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                if page_text:
                    text += page_text + "\n"
                else:
                    self.logger.warning(f"No text extracted from page {page_num + 1}")
            doc.close()
            
            text = text.strip()
            if not text:
                self.logger.warning("No text extracted from PDF using PyMuPDF, trying PyPDF fallback")
                raise Exception("No text extracted")
            
            self.logger.info(f"Successfully extracted {len(text)} characters from PDF using PyMuPDF")
            return text
            
        except Exception as e:
            self.logger.warning(f"PyMuPDF failed: {e}, trying PyPDF fallback")
            try:
                # Fallback to PyPDF
                with open(file_path, 'rb') as file:
                    reader = PdfReader(file)
                    text = ""
                    for page_num, page in enumerate(reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                        else:
                            self.logger.warning(f"No text extracted from page {page_num + 1} using PyPDF")
                    
                    text = text.strip()
                    if not text:
                        self.logger.error("No text could be extracted from PDF using either method")
                        return "PDF document (no text content extracted)"
                    
                    self.logger.info(f"Successfully extracted {len(text)} characters from PDF using PyPDF")
                    return text
                    
            except Exception as fallback_error:
                self.logger.error(f"Both PDF extraction methods failed: {fallback_error}")
                return "PDF document (text extraction failed)"
    
    def _process_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        doc = DocxDocument(str(file_path))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    
    def _process_text(self, file_path: Path) -> str:
        """Read text file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    
    def _process_json(self, file_path: Path) -> str:
        """Process JSON file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
            return json.dumps(data, indent=2)
    
    def _process_xml(self, file_path: Path) -> str:
        """Process XML file."""
        tree = ET.parse(file_path)
        root = tree.getroot()
        return ET.tostring(root, encoding='unicode')
    
    def _process_csv(self, file_path: Path) -> str:
        """Process CSV file."""
        df = pd.read_csv(file_path)
        return df.to_string(index=False)
    
    def _process_excel(self, file_path: Path) -> str:
        """Process Excel file."""
        df = pd.read_excel(file_path)
        return df.to_string(index=False) 